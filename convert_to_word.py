import os
import sys
import subprocess

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    print("Installing python-docx...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

MD_PATH = r"C:\Users\soman\Desktop\a\Detailed_Project_Report.md"
DOCX_PATH = r"C:\Users\soman\Desktop\a\Quantum_Exchange_College_Report.docx"

try:
    with open(MD_PATH, "r", encoding="utf-8") as f:
        lines = f.readlines()
        
    in_code_block = False
    for line in lines:
        if not in_code_block:
            line_stripped = line.strip()
        else:
            line_stripped = line.strip('\r\n')
            
        if line_stripped.startswith("```"):
            in_code_block = not in_code_block
            continue
            
        if not in_code_block:
            if line_stripped.startswith("# "):
                doc.add_heading(line_stripped[2:], level=1)
            elif line_stripped.startswith("## "):
                doc.add_heading(line_stripped[3:], level=2)
            elif line_stripped.startswith("### "):
                doc.add_heading(line_stripped[4:], level=3)
            elif line_stripped.startswith("#### "):
                doc.add_heading(line_stripped[5:], level=4)
            elif line_stripped == "" or line_stripped == "---":
                continue
            else:
                doc.add_paragraph(line_stripped)
        else:
            p = doc.add_paragraph(line_stripped)
            p.style = doc.styles['No Spacing']
            for run in p.runs:
                run.font.name = 'Courier New'

    doc.save(DOCX_PATH)
    print(f"SUCCESS: {DOCX_PATH} has been perfectly generated in your folder!")
except Exception as e:
    print(f"Error: {e}")
